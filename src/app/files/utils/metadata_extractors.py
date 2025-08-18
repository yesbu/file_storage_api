from pypdf import PdfReader
from docx import Document
import subprocess, tempfile, os
from datetime import datetime
def extract_pdf_meta(stream: bytes) -> dict:
    reader = PdfReader(stream)
    info = reader.metadata or {}
    pages = len(reader.pages)
    def _get(x):
        try:
            v = info.get(x)
            return str(v) if v is not None else None
        except Exception:
            return None
    return {"type": "pdf", "pages": pages, "author": _get("/Author"), "title": _get("/Title"), "created": _get("/CreationDate"), "producer": _get("/Producer") or _get("/Creator")}
def _convert_doc_to_docx(input_path: str) -> str:
    out_dir = os.path.dirname(input_path)
    subprocess.run(["libreoffice", "--headless", "--convert-to", "docx", "--outdir", out_dir, input_path], check=True)
    return input_path.rsplit(".", 1)[0] + ".docx"
def extract_docx_meta_from_path(path: str) -> dict:
    doc = Document(path)
    cp = doc.core_properties
    paragraphs = len(doc.paragraphs)
    tables = len(doc.tables)
    created = cp.created if cp.created else None
    if isinstance(created, datetime):
        created = created.isoformat()
    return {"type": "docx", "paragraphs": paragraphs, "tables": tables, "title": cp.title, "author": cp.author, "created": created}
def extract_office_meta(stream: bytes, filename: str) -> dict:
    suffix = filename.lower().split(".")[-1]
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, filename)
        with open(in_path, "wb") as f:
            f.write(stream)
        if suffix == "doc":
            docx_path = _convert_doc_to_docx(in_path)
            return extract_docx_meta_from_path(docx_path)
        else:
            return extract_docx_meta_from_path(in_path)

