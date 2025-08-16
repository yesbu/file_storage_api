from celery import Celery
from sqlalchemy.orm import Session
from . import models
from .database import SessionLocal
from .minio_client import minio_client  
import PyPDF2
from docx import Document
from io import BytesIO

app = Celery('tasks', broker='redis://redis:6379/0')

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.task(bind=True)  
def extract_metadata(self, file_id: int, content_type: str):
    db = SessionLocal()
    try:
        file = db.query(models.File).filter(models.File.id == file_id).first()
        if not file:
            return
        
        metadata = {}
        
        if content_type == "application/pdf":
            file_obj = minio_client.get_file(file.path)
            with BytesIO(file_obj.read()) as f:
                pdf = PyPDF2.PdfReader(f)
                metadata.update({
                    "pages": len(pdf.pages),
                    "author": pdf.metadata.get("/Author"),
                    "title": pdf.metadata.get("/Title")
                })
        
        elif content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            file_obj = minio_client.get_file(file.path)
            with BytesIO(file_obj.read()) as f:
                doc = Document(f)
                metadata.update({
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables)
                })
        
        
        if metadata: 
            db_metadata = models.FileMetadata(
                file_id=file.id,
                **metadata
            )
            db.add(db_metadata)
            db.commit()
        
    except Exception as e:
        db.rollback()
       
        raise self.retry(exc=e, countdown=60)
    finally:
        db.close()